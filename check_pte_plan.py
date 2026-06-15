#!/usr/bin/env python3
"""Check PTE plan files for garbled text issues."""
import sys

docx_path = "C:/Users/ikun/Desktop/PTE_4Week_Plan.docx"
pdf_path = "C:/Users/ikun/Desktop/PTE_4Week_Plan.pdf"

# Check DOCX
print("=" * 60)
print("CHECKING DOCX...")
print("=" * 60)
try:
    from docx import Document
    doc = Document(docx_path)
    para_count = sum(1 for p in doc.paragraphs if p.text.strip())
    print(f"Paragraphs with text: {para_count}")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            # Check for garbled content
            has_garbled = any(ord(c) > 127 and not (
                '\u4e00' <= c <= '\u9fff' or  # CJK
                '\u3000' <= c <= '\u303f' or  # CJK punctuation
                '\uff00' <= c <= '\uffef' or  # Fullwidth
                c in '，。、；：？！""''【】（）—…·'
            ) for c in t)
            marker = " [GARBLED?]" if has_garbled else ""
            print(f"  [{i}]{marker} {t[:150]}")
    
    # Check tables
    print(f"\nTables: {len(doc.tables)}")
    for ti, table in enumerate(doc.tables):
        print(f"  Table {ti}: {len(table.rows)}x{len(table.columns)}")
        for ri, row in enumerate(table.rows):
            cells_text = [c.text[:40].replace('\n', '|') for c in row.cells]
            print(f"    Row {ri}: {cells_text}")
except ImportError:
    print("python-docx not installed")
    # Manual zip inspection
    import zipfile
    import xml.etree.ElementTree as ET
    with zipfile.ZipFile(docx_path) as z:
        for name in z.namelist():
            if name.startswith('word/document') or name.startswith('word/header') or name.startswith('word/footer'):
                print(f"\n--- {name} ---")
                xml_content = z.read(name)
                # Try to extract text
                root = ET.fromstring(xml_content)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                texts = root.findall('.//w:t', ns)
                full_text = ''.join(t.text or '' for t in texts)
                print(full_text[:1000])
except Exception as e:
    print(f"Error: {e}")

print("\n" + "=" * 60)
print("CHECKING PDF...")
print("=" * 60)
try:
    from pdfminer.high_level import extract_text
    text = extract_text(pdf_path)
    garbled_chars = sum(1 for c in text if ord(c) > 127 and not (
        '\u4e00' <= c <= '\u9fff' or
        '\u3000' <= c <= '\u303f' or
        '\uff00' <= c <= '\uffef' or
        c in '，。、；：？！""''【】（）—…·'
    ))
    print(f"Total chars: {len(text)}, garbled chars: {garbled_chars}")
    print(text[:2000])
except ImportError:
    print("pdfminer not installed, trying pypdf...")
    try:
        import pypdf
        reader = pypdf.PdfReader(pdf_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            print(f"\n--- Page {i+1} ---")
            print(text[:500])
    except ImportError:
        print("No PDF library available")
except Exception as e:
    print(f"Error: {e}")
